import os
import json
import google.generativeai as genai
import openpyxl
import logging
import argparse
from openpyxl.styles import Alignment
from openpyxl.utils.datetime import from_excel
from dotenv import load_dotenv
import subprocess
import concurrent.futures
import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import sys
from pathlib import Path
import time
import google.api_core.exceptions
from docx import Document
import datetime
import xml.parsers.expat
import webbrowser

# ユーザーディレクトリのDocumentsフォルダのパスを取得
documents_path = Path.home() / "Documents"
log_file_path = documents_path / "app_log.txt"

# ログファイルの設定
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file_path, encoding='utf-8'),  # エンコーディングを指定
        logging.StreamHandler()
    ]
)

def get_current_dir():
    # この関数は、現在のスクリプトがどこにあるかを教えてくれます。
    if getattr(sys, 'frozen', False):
        # もしプログラムがPyInstallerでパッケージ化されているなら
        return Path(sys.executable).resolve().parent  # 修正
    else:
        # そうでないなら、開発中のフォルダを使います
        return Path(__file__).resolve().parent

settings_path = os.path.join(get_current_dir(), 'settings.json')
print(f"Settings path: {settings_path}")
print(f"File exists: {os.path.exists(settings_path)}")

# 現在のスクリプトのディレクトリを取得
current_dir = Path(__file__).resolve().parent

# 環境変数の読み込み
load_dotenv(current_dir / 'env_variables.env')  # ファイル名を英語に変更

# プロジェクトディレクトリの設定
project_dir = os.path.dirname(os.path.abspath(__file__))

# APIキーの設定
API_KEYS = [os.getenv(f'GEMINI_API_KEY_{i}') for i in range(1, 11)]  # 10個のAPIキーを取得

# 処理済みファイルのログファイル
PROCESSED_FILES_LOG = os.path.join(current_dir, 'processed_files.json')

def load_processed_files():
    # この関数は、すでに処理したファイルのリストを読み込みます。
    if os.path.exists(PROCESSED_FILES_LOG):
        # もしログファイルが存在するなら
        with open(PROCESSED_FILES_LOG, 'r', encoding='utf-8') as f:
            return json.load(f)  # ファイルを開いて内容を読み込みます
    return {}  # ファイルがなけば空のリストを返します

def save_processed_files(processed_files):
    # この関数は、処理したファイルのリストを保存します。
    with open(PROCESSED_FILES_LOG, 'w', encoding='utf-8') as f:
        json.dump(processed_files, f, indent=2, ensure_ascii=False)  # ファイルに書き込みます

def get_unprocessed_audio_files():
    # この関数は、まだ処理していない音声ファイルを探します。
    processed_files = load_processed_files()  # すでに処理したファイルを取得
    audio_files = [f for f in os.listdir(current_dir) if f.endswith('.mp3')]
    # フォルダ内のすべての.mp3ファイルをリストにします
    return [f for f in audio_files if f not in processed_files]
    # まだ処理していないファイルだけを返します

def create_extraction_prompt(text):
    # この関数は、会議の内容から情報を抽出するための指示を作ります。
    return f"""
    この文章はとある会議の内容です。
    以下の文章から、次の項目を抽出してください：
    1. 議題①
    2. 議題①の要約
    3. 議題②
    4. 議題②の要約
    5. 議題③
    6. 議題③の要約
    7. 議題④
    8. 議④の要約
    9. 議題⑤
    10. 議題⑤の要約
    11. 議題⑥
    12. 議題⑥の要約
    13. 議題
    14. 議題⑦の要約
    15. 議題⑧
    16. 議題⑧の要約
    17. 議題⑨
    18. 議題⑨の要約
    19. 議題⑩
    20. 議題⑩の要約
    21. 議題⑪
    22. 議題⑪の要約
    23. 議題⑫
    24. 議題⑫の要約
    25. 議題⑬
    26. 議題⑬の要約
    27. 議題⑭
    28. 議題⑭の要約
    29. 議題⑮
    30. 議題⑮の要約
    31. 議題⑯
    32. 議題⑯の要約
    33. 議題⑰
    34. 議題⑰の要約
    35. 議題⑱
    36. 議題⑱の要約
    37. 議題⑲
    38. 議題⑲の要約
    39. 議題⑳
    40. 議題⑳の要約

    抽出する際は、必ず以下の形式で出力してください：
    議題①: [議題の内容]
    議題①の要約: [要約内容]

    議題②: [議題の内容]
    議題②の要約: [要約内容]

    議題③: [議題の内容]
    議題③の要約: [要約内容]

    ...

    議題⑳: [議題の内容]
    議題⑳の要約: [要約内容]

    注意事項:
    - 各議題とその要約を必ず上記の形式で出力してください。
    - 議題が20個未満の場合は、存在する議題のみを抽出してください。
    - 要約は簡潔かつ具体的にしてください。
    - 議題の番号（①、②など）は必ず付けてください。
    - 各行は必ず「議題○:」または「議題○の要約:」で始まるようにしてください。
    - 議題や要約の前に「*」や「**」などの記号を付けないでください。
    - 議題というのはあくまで表現の一つであり、会話内容が議事録形式で記されていれば構いません。インタビューの文章等からも適切に議題を抽出してい。
    - インタビーのような文章であっても、適切に議題を抽出してください。

    文章:
    {text}
    """

def get_ffmpeg_path():
    if getattr(sys, 'frozen', False):
        return os.path.join(sys._MEIPASS, 'ffmpeg.exe')
    return os.path.join(os.path.dirname(__file__), 'ffmpeg.exe')

def get_ffprobe_path():
    if getattr(sys, 'frozen', False):
        return os.path.join(sys._MEIPASS, 'ffprobe.exe')
    return os.path.join(os.path.dirname(__file__), 'ffprobe.exe')

def split_audio_file(audio_file_path, num_parts):
    """音声ファイルを指定された数の部分に重なりを持たせて分割する関数"""
    # この関数は、長い音声ファイルを小さな部分に分けます。
    # 分けた部分は少し重なりを持つので、途切れないようになっています。

    file_size = os.path.getsize(audio_file_path)  # ファイルの大きさを調べます
    duration = get_audio_duration(audio_file_path)  # 音声ファイルの長さを取得します
    part_duration = duration / num_parts  # 各部分の長さを計算します
    overlap_duration = part_duration * 0.1  # 10%の重なりを持たせます

    parts = []  # 分割した声フイルのリストを作ります
    for i in range(num_parts):
        # 部分の開始時間を計算します
        start_time = max(0, i * part_duration - (overlap_duration if i > 0 else 0))
        # 新しい音声ファイルの名前を決めます
        part_file = f"{audio_file_path}_part{i+1}.mp3"  # 拡張子をmp3のままにします

        # 音声ファイルの種類に応じて、分割の法を変えます
        if audio_file_path.endswith('.mp3'):
            # MP3ファイルの場合の分割方法
            command = [
                str(get_ffmpeg_path()),  # ffmpegといソフトウェアのパスを取得します
                '-y',  # 同じ名前のファイルがあれば上書きします
                '-i', audio_file_path,  # 元の音声ファイルを指定します
                '-ss', str(start_time),  # 開始時間を指定します
                '-t', str(part_duration + (overlap_duration if i < num_parts - 1 else 0)),  # 部分の長さを指定します
                '-c', 'copy',  # 音声をそのままコピーします（音質を変えません）
                part_file  # 新しい音声ファイルの名前を指定します
            ]
        elif audio_file_path.endswith('.m4a'):
            # M4Aファイルの場合の分方法
            part_file = f"{audio_file_path}_part{i+1}.m4a"  # 拡m4aまま
            command = [
                str(get_ffmpeg_path()),
                '-y',
                '-i', audio_file_path,
                '-ss', str(start_time),
                '-t', str(part_duration + (overlap_duration if i < num_parts - 1 else 0)),
                '-c', 'copy',
                part_file
            ]
        elif audio_file_path.endswith('.wav'):
            # WAVファイルの場合の分割方法
            part_file = f"{audio_file_path}_part{i+1}.wav"  # 拡張子をwavのままにしま
            command = [
                str(get_ffmpeg_path()),
                '-y',
                '-i', audio_file_path,
                '-ss', str(start_time),
                '-t', str(part_duration + (overlap_duration if i < num_parts - 1 else 0)),
                '-c', 'pcm_s16le',  # WAV用の音声形式を指定します
                part_file
            ]

        # 音声ファイルを実際に分割します
        result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if result.returncode != 0:
            # エラーが起きた場合は記録します
            logging.error(f"FFmpegエラー: {result.stderr}")
        parts.append(part_file)  # 分割したファイルをリストに追加します

    return parts  # 割したファイルのリストを返しす

def get_audio_duration(audio_file_path):
    """音声ファイルの長さを取得する関数"""
    # この関数は、音声ファイルの再生時間（長さ）を秒単位で取得します

    # ffprobeというツールを使うためのコマンドを準備します
    ffprobe_path = get_ffprobe_path()
    command = [ffprobe_path, "-v", "error", "-show_entries", "format=duration", "-of", "default=noprint_wrappers=1:nokey=1", audio_file_path]

    # 準備したコマンドを実行し、結果を取得します
    result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

    # 結果から音声ファイルの長さ（秒）を取り出し、小数点の数値として返します
    return float(result.stdout.strip())

# グローバル変数の定義
transcription_prompt = ""

def load_prompt_from_settings():
    """settings.jsonからプロンプトを読み込む関数"""
    settings_path = get_settings_path()
    logging.info(f"Settings path: {settings_path}")  # 追加: パスをログに出力
    if os.path.exists(settings_path):
        logging.info("settings.jsonが見つかりました。")  # 追加: ファイ存在確認
        with open(settings_path, 'r', encoding='utf-8') as f:
            try:
                settings = json.load(f)
                logging.info("settings.jsonを正常に読み込みました。")  # 追加: 読み込み成功
                return settings.get('transcription_prompt', '')  # デフォルト値を空文字に変更
            except json.JSONDecodeError as e:
                logging.error(f"JSONデコーエラー: {str(e)}")  # 追加: JSONデコードエラー
    else:
        logging.error("settings.jsonが見つかりません。")  # 追加: ファイルが見つからない場合
    return ''  # ファイルが存在しない場合も空文字を返す

def transcribe_audio_with_key(audio_file, api_key, retries=3):
    """指定されたAPIキーを使用て音声ファイルを文字起こしする関数"""
    # この関数は、音声ファイルをテキストに変換します

    # プロンプトをログに出力（1回だけ）
    if transcription_prompt:
        # もし文字起こしの指示（プロンプト）があれば、それをログに記録します
        logging.info(f"今回は以下のプロンプトで文字起こしします:\n{transcription_prompt}")
    else:
        # プロンプトがない場合はエラーを記録して、関数を終了します
        logging.error("プロンプトが取得できませんでした。")
        return None

    # 指定さた回数（デフォルトは3回）まで文字起こしを試みます
    for attempt in range(retries):
        try:
            # 音声ファイルを開いてデータを読み込みます
            with open(audio_file, 'rb') as audio:
                audio_data = audio.read()

            # Geminiモデルを設定します
            model = genai.GenerativeModel('gemini-1.5-pro')
            genai.configure(api_key=api_key)

            # モデルを使って音声データを文字に起こします
            response = model.generate_content(
                [
                    transcription_prompt,
                    {"mime_type": "audio/mp3", "data": audio_data}
                ]
            )

            # 文字起こしが成功したかチェックします
            if hasattr(response, 'text'):
                # 成功した場合、ログに記録して結果を返します
                logging.info(f"{audio_file}の文字起こしが成功しました。")
                return response.text
            else:
                # テキストが含まれていない場合はエラーを記録します
                logging.error(f"文起こし失敗: {audio_file} - レスポンスにテキストが含まれていません。")
        except google.api_core.exceptions.ResourceExhausted:
            # APIの利用制限に達した場合のエラーを録します
            logging.error(f"文字起こし失敗: {audio_file} - 429 Resource has been exhausted (e.g. check quota).")
        except Exception as e:
            # その他のエラーが発生し、エ内容を記録します
            logging.error(f"文字起こし失敗: {audio_file} - {str(e)}")
        
        # リトライが可能な場合は、次の試行を行います
        if attempt < retries - 1:
            logging.info(f"リトライを試みます ({attempt + 2}/{retries})")
            time.sleep(30)  # 1分待ってから次の試行を行います
        else:
            # すべての試行が失敗した場合、最終的なエラーを記録します
            logging.error(f"{audio_file}の文字起こしが{retries}回失敗しました。")
    
    # すべての試行が失敗した場合はNoneを返します
    return None

def extract_information(text, api_key):
    # この関数は、テキストから重要な情報を抽出します

    # テキストの空白を整理します
    cleaned_text = " ".join(text.split())

    # APIキーがない場合はエラーを記録して終了します
    if not api_key:
        logging.error("情報抽出に使用するAPIキーが設定されていません。")
        return

    # APIキーを設定して、AIモデルを準備します
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-pro')
    
    # 情報抽出のための指示文を作ります
    prompt = create_extraction_prompt(cleaned_text)

    try:
        # 情報抽出を開始します
        logging.info("情報抽出を開始します。")
        # AIモデルに指示を送り、結果を受け取ります
        response = model.generate_content(prompt)
        # 結果のテキストから余分な空白を取り除きます
        extracted_text = response.text.strip()
        # 抽出結果を記録します
        logging.info(f"抽出結果全体: {extracted_text}")
        # 抽出したテキストを返します
        return extracted_text
    except Exception as e:
        # エラーが起きた場合、詳細を記録して再度エラーを発生させます
        logging.exception(f"情報抽出中にエラーが発生しました: {str(e)}")
        raise

def create_excel(extracted_info, output_file):
    # 新しいExcelワークブックを作成します
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "議事録"

    # 列の幅を設定します
    ws.column_dimensions['A'].width = 20  # A列（議題）の幅を20に設定
    ws.column_dimensions['B'].width = 80  # B列（内容）の幅を80に設定

    # 会議詳細情報を追加します
    meeting_details = [
        "会議名",
        "日時",
        "場所",
        "参加者",
        "欠席者"
    ]

    # 会議詳細情報をExcelに書き込みます
    for i, detail in enumerate(meeting_details, start=1):
        ws.cell(row=i, column=1, value=detail)  # A列に項目名を書き込み
        ws.cell(row=i, column=1).font = openpyxl.styles.Font(bold=True)  # 太字に設定
        ws.cell(row=i, column=1).fill = openpyxl.styles.PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")  # 背景色を設定

    row = 6  # 会議詳細情報の後から議題の書き込みを開始します

    # 抽出された情報を行ごとに分割します
    lines = extracted_info.split('\n')
    current_topic = ""
    current_summary = ""

    # 各行を処理して議題と要約を抽出します
    for line in lines:
        line = line.strip()
        if line.startswith("議題"):
            if current_topic and current_summary:
                # 前の議題を書き込みます
                ws.cell(row=row, column=1, value=current_topic)
                cell = ws.cell(row=row, column=2, value=current_summary)
                cell.alignment = Alignment(wrap_text=True)  # テキストを折り返して表示
                row += 1
            parts = line.split(':', 1)
            if len(parts) == 2:
                current_topic = parts[0].strip()
                current_summary = parts[1].strip()
            else:
                current_topic = line
                current_summary = ""
        elif "の要約" in line:
            if current_topic and "の要約:" in line:
                current_summary = line.split("の要約:", 1)[1].strip()
        elif current_summary:
            current_summary += " " + line.strip()

    # 最後の議題を書き込みます
    if current_topic and current_summary:
        ws.cell(row=row, column=1, value=current_topic)
        cell = ws.cell(row=row, column=2, value=current_summary)
        cell.alignment = Alignment(wrap_text=True)

    # セルのスタイルを設定します
    for row in ws['A1:B'+str(ws.max_row)]:
        for cell in row:
            # すべてのセルに細い枠線を設定
            cell.border = openpyxl.styles.Border(left=openpyxl.styles.Side(style='thin'), 
                                                 right=openpyxl.styles.Side(style='thin'), 
                                                 top=openpyxl.styles.Side(style='thin'), 
                                                 bottom=openpyxl.styles.Side(style='thin'))
            if cell.column == 1:  # A列（議題）のセルの場合
                cell.font = openpyxl.styles.Font(bold=True)  # 太字に設定
                cell.fill = openpyxl.styles.PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")  # 背景色を設定
            elif cell.column == 2:  # B列（内容）のセルの場合
                cell.alignment = Alignment(wrap_text=True)  # テキストを折り返して表示

    # B列の幅を内容に合わせて自動調整します
    for column_cells in ws.columns:
        length = max(len(str(cell.value)) for cell in column_cells)
        if column_cells[0].column_letter == 'B':
            ws.column_dimensions[column_cells[0].column_letter].width = min(100, max(80, length))

    # Excelファイルを保存します
    try:
        wb.save(output_file)
        logging.info(f"Excelファイルが正常に作成されました: {output_file}")
    except PermissionError:
        logging.error(f"Excelファイルの保存に失敗しました。書き込み権限がありません: {output_file}")
    except Exception as e:
        logging.error(f"Excelファイルの保存中にエラーが発生しました: {str(e)}")

def load_output_directory():
    """settings.jsonから出力先ディレクトリを読み込む関数"""
    settings_path = get_settings_path()
    if os.path.exists(settings_path):
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
            return settings.get('output_directory', os.path.join(Path.home(), 'Documents'))
    return os.path.join(Path.home(), 'Documents')

def process_audio_file(audio_file_path, processed_files):
    try:
        audio_file_name = os.path.basename(audio_file_path)
        file_size = os.path.getsize(audio_file_path)
        logging.info(f"{audio_file_name}の処理を開始します。ファイルサイズ: {file_size / (1024 * 1024):.2f}MB")

        # APIキーをロード
        api_keys = load_api_keys()
        if not api_keys:
            logging.error("APIキーがロードされていません。処理を中止します。")
            return False

        # 音声ファイルを分割する数を決定
        num_parts = len(api_keys)  # APIキーの数に応じて分割数を決定

        transcribed_texts = [None] * num_parts  # インデックスに基づいて配置するリスト

        audio_parts = split_audio_file(audio_file_path, num_parts)

        with concurrent.futures.ThreadPoolExecutor() as executor:
            future_to_index = {executor.submit(transcribe_audio_with_key, part, api_keys[i]): i for i, part in enumerate(audio_parts)}
            failed_parts = []
            successful_api_keys = []  # 成功したAPIキーを記録るリスト
            for future in concurrent.futures.as_completed(future_to_index):
                index = future_to_index[future]
                part = audio_parts[index]
                result = future.result()
                if result:
                    transcribed_texts[index] = result
                    logging.info(f"{part}の処理が成功しました。")
                    if api_keys[index] not in successful_api_keys:
                        successful_api_keys.append(api_keys[index])  # 成功したAPIキーを記録
                else:
                    logging.error(f"{part}の処理が失敗しました。")
                    failed_parts.append((index, part))

        # 失敗したパートのリトライ
        if failed_parts:
            logging.info("失敗したファイルのリトライを1分後に開始します。")
            time.sleep(60)
            for index, part in failed_parts:
                result = transcribe_audio_with_key(part, api_keys[0])
                if result:
                    transcribed_texts[index] = result
                    logging.info(f"{part}のリトライが成功しました。")
                    if api_keys[0] not in successful_api_keys:
                        successful_api_keys.append(api_keys[0])  # リトライで成功しAPIキーを記録
                else:
                    logging.error(f"{part}のリトライが失敗しました。")

        # 分割されたファイルを削除
        for part in audio_parts:
            os.remove(part)
        logging.info(f"{audio_file_name}の分割されたファイルを削除しました。")

        # 文字起こし結果を結合（Noneを除外）
        combined_text = "\n".join(filter(None, transcribed_texts))
        # 余分な空白を取り除く
        cleaned_combined_text = " ".join(combined_text.split())
        logging.info(f"{audio_file_name}字起こしが完了しました。情報を抽出します。")

        # 文字起こし結果をWordファイルに保存
        try:
            output_directory = load_output_directory()
            word_output_file = os.path.join(output_directory, f"{os.path.splitext(audio_file_name)[0]}_文字起こし.docx")
            doc = Document()
            doc.add_paragraph(cleaned_combined_text)
            doc.save(word_output_file)
            logging.info(f"文字起こし結果がWordファイルに保存されました: {word_output_file}")
        except Exception as e:
            logging.error(f"文字起こし結果のWordファイル保存中にエラーが発生しました: {str(e)}")
            return False

        # 15秒のバッファを持たせる
        time.sleep(15)

        # 成功したAPIキー使って情報抽出を試みる
        for api_key in successful_api_keys:
            try:
                extracted_info = extract_information(cleaned_combined_text, api_key)
                if extracted_info:
                    output_file = os.path.join(output_directory, f"{os.path.splitext(audio_file_name)[0]}_抽出結果.xlsx")
                    create_excel(extracted_info, output_file)
                    processed_files[audio_file_name] = output_file
                    break  # 成功したらループを抜ける
            except google.api_core.exceptions.ResourceExhausted:
                logging.error(f"{api_key}での情報抽出が失敗しました。次のAPIキーを試します。")
        else:
            logging.error(f"{audio_file_name}の情報抽出に失敗しました。")

        return True
    except Exception as e:
        logging.exception(f"{audio_file_path}の処理中にエラーが発生しました: {str(e)}")
        return False

def extract_info_from_xlsx(file_path):
    wb = openpyxl.load_workbook(file_path)
    sheet = wb.active
    data = {
        '会議名': sheet['B1'].value or '',
        '日時': convert_excel_date(sheet['B2'].value),
        '場所': sheet['B3'].value or '',
        '参加者': sheet['B4'].value or '',
        '欠席者': sheet['B5'].value or '',
    }
    for i in range(1, 21):  # 議題①から⑩まで
        data[f'議題{chr(0x2460 + i - 1)}'] = sheet[f'B{5+i*2-1}'].value or ''
        data[f'議題{chr(0x2460 + i - 1)}の要約'] = sheet[f'B{5+i*2}'].value or ''
    
    print("抽出されたデータ:")
    for key, value in data.items():
        print(f"{key}: {value}")
    
    return data

def convert_excel_date(value):
    if isinstance(value, (int, float)):
        return from_excel(value).strftime('%Y-%m-%d')
    return value

def create_minutes_from_template(data, template_path):
    try:
        if getattr(sys, 'frozen', False):
            # PyInstallerでパッケージ化されている場合
            template_path = os.path.join(sys._MEIPASS, 'template.docx')
        else:
            # 通常のPython実行の場合
            template_path = os.path.join(get_current_dir(), 'template.docx')

        print(f"Using template path: {template_path}")  # デバッグ用
        print(f"Template exists: {os.path.exists(template_path)}")  # デバッグ用

        doc = Document(template_path)
        
        for paragraph in doc.paragraphs:
            # 会議名、日時、場所、参加者、欠席者の置き換え
            for key, value in data.items():
                placeholder = f'「{key}」'
                if placeholder in paragraph.text:
                    old_text = paragraph.text
                    new_text = paragraph.text.replace(placeholder, str(value) if value is not None else '')
                    paragraph.text = new_text
                    print(f"置換: '{old_text}' -> '{new_text}'")

            # 議題と要約の置き換え
            for i in range(1,21):
                topic_key = f'議題{chr(0x2460 + i - 1)}'
                topic_content = data.get(topic_key, '')
                summary_key = f'議題{chr(0x2460 + i - 1)}の要約'
                summary_content = data.get(summary_key, '')

                # 議題の名称を置き換え
                topic_placeholder = f'「{topic_key}」'
                if topic_placeholder in paragraph.text:
                    old_text = paragraph.text
                    new_text = paragraph.text.replace(topic_placeholder, topic_content)
                    paragraph.text = new_text
                    print(f"議題名置換: '{old_text}' -> '{new_text}'")

                # 要約の置き換え
                summary_placeholder = f'「{summary_key}」'
                if summary_placeholder in paragraph.text:
                    old_text = paragraph.text
                    new_text = paragraph.text.replace(summary_placeholder, summary_content)
                    paragraph.text = new_text
                    print(f"要約置換: '{old_text}' -> '{new_text}'")

        return doc
    except Exception as e:
        logging.error(f"テンプレート処理中にエラーが発生: {str(e)}")
        raise

def create_minutes(xlsx_path, template_path, output_path):
    try:
        data = extract_info_from_xlsx(xlsx_path)
        doc = create_minutes_from_template(data, template_path)
        doc.save(output_path)
        print(f"議事録が作成されました: {output_path}")
        return True
    except Exception as e:
        logging.error(f"議事録の作成中にエラーが発生しました: {str(e)}")
        print(f"エラーが発生しました: {str(e)}")
        return False
    
# グローバル変数の定義
selected_file = None
file_label = None
excel_file_label = None
uploading_label = None
elapsed_time_label = None
estimated_time_label = None  # 追加
root = None
processing_done = False  # 処理が完了したかどうかを示すフラグ
start_time = None  # 処理開始時刻を保持
selected_file_name = ""  # 選択したファイル名を保持
estimated_time_text = ""  # 想定処理時間を保持

def show_main_menu():
    global root, file_label, excel_file_label, uploading_label, elapsed_time_label, estimated_time_label, selected_file, selected_file_name, estimated_time_text,transcription_prompt, processing_done, start_time
    for widget in root.winfo_children():
        widget.destroy()

    root.title("爆速議事録")
    root.geometry("900x600")  # ウィンドウサイズを大きくする
    root.resizable(False, False)  # ウィンドウのサイズを固定

    settings = load_settings()
    transcription_prompt = settings.get('transcription_prompt', '')
    output_directory = settings.get('output_directory', '')
    api_keys = [settings.get('gemini_api_keys', {}).get(f'GEMINI_API_KEY_{i}', '') for i in range(1, 11)]

    # 設定の読み込み状態をログに記録
    logging.info(f"設定を読み込みました:")
    logging.info(f"- プロンプト: {'設定済み' if transcription_prompt else '未設定'}")
    logging.info(f"- 出力先ディレクトリ: {output_directory if output_directory else '��設定'}")
    logging.info(f"- APIキー: {sum(1 for key in api_keys if key) }個設定済み")

    # タイトルラベル
    title_label = tk.Label(root, text="⚡️爆速議事録", font=("Yu Gothic", 24, "bold"))
    title_label.pack(pady=20)

    # 設定ボタンを右上に配置
    settings_button = tk.Button(root, text="設定", command=show_settings, width=8, height=1)
    settings_button.place(x=800, y=20)

    # 使い方ボタンを設定ボタンの下に配置
    usage_button = tk.Button(root, text="使い方", command=show_usage, width=8, height=1)
    usage_button.place(x=800, y=60)

    # メインフレーム
    main_frame = tk.Frame(root)
    main_frame.pack(expand=True, fill="both", padx=20, pady=20)

    # 音声ファイル処理フレーム
    audio_frame = tk.LabelFrame(main_frame, text="音声ファイル処理", font=("Yu Gothic", 12, "bold"), padx=10, pady=10)
    audio_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")

    audio_button = tk.Button(audio_frame, text="音声ファイルを選択する", command=upload_audio_file, width=25)
    audio_button.pack(pady=10)

    file_label = tk.Label(audio_frame, text=f"選択したファイル: {selected_file_name}", wraplength=300, justify="center")
    file_label.pack(pady=10)

    # 音声ファイルを処理するボタン
    process_audio_button = tk.Button(audio_frame, text="音声ファイルを処理する", command=complete_audio_upload, width=25)
    process_audio_button.pack(pady=10)  

    # 想定処理時間を表示するラベル
    estimated_time_label = tk.Label(audio_frame, text="", font=("Arial", 12))
    estimated_time_label.pack(pady=10)

    # 経過時間を表示するラベル
    uploading_label = tk.Label(audio_frame, text="", font=("Arial", 12))
    uploading_label.pack(pady=10)

    # 処理が進行中の場合、選択ファイルと想定時間を表示
    if not processing_done and selected_file and start_time:
        file_label.config(text=f"選択したファイル: {os.path.basename(selected_file)}")
        estimated_time_label.config(text=f"想定処理時間: {estimated_time_text}")
        
        def update_elapsed_time(start_time=start_time):
            if processing_done:
                uploading_label.config(text="")  # 経過時間表示をクリア
                return
                
            if start_time:
                try:
                    elapsed_time = int(time.time() - start_time)
                    minutes, seconds = divmod(elapsed_time, 60)
                    uploading_label.config(text=f"経過時間: {minutes}分{seconds}秒")
                    if not processing_done:
                        root.after(1000, lambda: update_elapsed_time(start_time))
                except Exception as e:
                    logging.error(f"経過時間の更新中にエラーが発生: {str(e)}")

        # メインスレッドで更新を開始
        root.after(0, update_elapsed_time)

    # Excelファイル処理フレーム
    excel_frame = tk.LabelFrame(main_frame, text="Excelファイル処理", font=("Yu Gothic", 12, "bold"), padx=10, pady=10)
    excel_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")

    excel_button = tk.Button(excel_frame, text="Excelファイルを選択する", command=upload_xlsx_file, width=25)
    excel_button.pack(pady=10)

    excel_file_label = tk.Label(excel_frame, text="選択したファイル: なし", wraplength=300, justify="center")
    excel_file_label.pack(pady=10)

    process_excel_button = tk.Button(excel_frame, text="Excelファイルを処理する", command=complete_xlsx_upload, width=25)
    process_excel_button.pack(pady=10)

    # グリッドの設定
    main_frame.grid_columnconfigure(0, weight=1)
    main_frame.grid_columnconfigure(1, weight=1)
    main_frame.grid_rowconfigure(0, weight=1)

def show_settings():
    global root, file_label, excel_file_label, uploading_label, elapsed_time_label, estimated_time_label, selected_file, selected_file_name, estimated_time_text,transcription_prompt, processing_done, start_time
    for widget in root.winfo_children():
        widget.destroy()

    root.title("設定")
    
    # メインフレーム
    main_frame = tk.Frame(root)
    main_frame.pack(expand=True, fill="both", padx=20, pady=20)
    
    # 戻るボタン
    back_button = tk.Button(root, text="戻る", command=show_main_menu, width=8, height=1)
    back_button.place(x=800, y=20)
    back_button.lift()

    # 上部のスペース
    tk.Label(main_frame, height=1).pack()
    
    # コンテンツフレームのレイアウトを変更
    content_frame = tk.Frame(main_frame)
    content_frame.pack(expand=True, fill="both")
    
    # 左側（プロンプト）のフレーム - 幅を60%に設定
    left_frame = tk.LabelFrame(content_frame, text="文字起こしプロンプト", font=("Yu Gothic", 12, "bold"))
    left_frame.pack(side="left", fill="both", expand=True, padx=10, pady=10)
    
    # プロンプトテキストボックスとスクロールバー
    prompt_frame = tk.Frame(left_frame)
    prompt_frame.pack(fill="both", expand=True, padx=5, pady=5)
    
    prompt_textbox = tk.Text(prompt_frame, wrap="word", font=("Yu Gothic", 10), height=20)
    prompt_scrollbar = tk.Scrollbar(prompt_frame, orient="vertical", command=prompt_textbox.yview)
    prompt_textbox.configure(yscrollcommand=prompt_scrollbar.set)
    
    # 保存されているプロンプトを読み込んで表示
    settings = load_settings()
    saved_prompt = settings.get('transcription_prompt', '')
    prompt_textbox.insert('1.0', saved_prompt)
    
    prompt_scrollbar.pack(side="right", fill="y")
    prompt_textbox.pack(side="left", fill="both", expand=True)
    
    # プロンプトの保存ボタン
    save_prompt_button = tk.Button(left_frame, text="保存", 
                                command=lambda: save_prompt_to_settings(prompt_textbox.get('1.0', 'end-1c')))
    save_prompt_button.pack(side="bottom", pady=5, padx=5, fill="x")
    
    right_frame = tk.Frame(content_frame)
    right_frame.pack(side="right", fill="both", padx=10, pady=10, expand=True)
    
    # 出力先ディレクトリフレーム
    directory_frame = tk.LabelFrame(right_frame, text="出力先ディレクトリ", font=("Yu Gothic", 12, "bold"))
    directory_frame.pack(fill="x", padx=5, pady=(0, 10))
    
    current_dir = load_output_directory()
    current_dir_label = tk.Label(directory_frame, text=f"現在の出力先:\n{current_dir}", 
                                wraplength=300, font=("Yu Gothic", 10))
    current_dir_label.pack(pady=5)
    
    directory_button = tk.Button(directory_frame, text="ディレクトリを指定する", 
                               command=lambda: select_directory(current_dir_label))
    directory_button.pack(fill="x", pady=5, padx=5)
    
    # APIキーフレーム
    api_key_frame = tk.LabelFrame(right_frame, text="Gemini APIキー", font=("Yu Gothic", 12, "bold"))
    api_key_frame.pack(fill="both", expand=True, padx=5, pady=(0, 5))
    
    # APIキー入力フレーム
    api_key_input_frame = tk.Frame(api_key_frame)
    api_key_input_frame.pack(fill="both", expand=True, padx=5, pady=5)
    
    # APIキー入力用テキストボックスとスクロールバー
    api_key_textbox = tk.Text(api_key_input_frame, wrap="word", font=("Yu Gothic", 8), height=3)
    api_key_scrollbar = tk.Scrollbar(api_key_input_frame, orient="vertical", command=api_key_textbox.yview)
    api_key_textbox.configure(yscrollcommand=api_key_scrollbar.set)
    
    # 保存されているAPIキーを読み込んで表示
    api_keys_text = get_api_keys_text()
    api_key_textbox.insert('1.0', api_keys_text)
    
    api_key_scrollbar.pack(side="right", fill="y")
    api_key_textbox.pack(side="left", fill="both", expand=True)
    
    # APIキーの保存ボタン
    save_api_key_button = tk.Button(api_key_frame, text="保存", 
                                  command=lambda: save_api_keys_to_settings(api_key_textbox.get('1.0', 'end-1c')))
    save_api_key_button.pack(side="bottom", fill="x", pady=5, padx=5)

def select_directory(label_widget):
    directory = filedialog.askdirectory()
    if directory:
        label_widget.config(text=f"選択されたディレクトリ:\n{directory}")
        save_output_directory_to_settings(directory.strip())

def main():
    global root, transcription_prompt, selected_file_name, estimated_time_text
    try:
        logging.info("プロンプトをロード中...")  # 追加: ロード開始ログ
        transcription_prompt = load_prompt_from_settings()  # プロンプトをロード
        logging.info(f"取得したプロンプト: {transcription_prompt}")  # プロンプトの内容をログに出力
        logging.info("プロンプトのロードが完了しました。")  # 追加: ロード完了ログ
        root = tk.Tk()
        root.title("ファイル処理ツール")
        root.geometry("500x300")

        show_main_menu()

        root.mainloop()
    except Exception as e:
        logging.exception("アプリケーションの実行中にエラーが発生しました。")
        messagebox.showerror("エラー", f"アプリケーションの実行中にエラーが発生しました:\n{str(e)}")
        logging.error(f"アプリケーションの起動時にエラーが発生しました: {str(e)}")  # エラーログを追加

def show_usage():
    for widget in root.winfo_children():
        widget.destroy()

    root.title("使い方")

    usage_label = tk.Label(root, text="使い方", font=("Yu Gothic", 16, "bold"))
    usage_label.pack(pady=(60, 0))  # 上に60ピクセルの余白を追加

    usage_text = "使い方は以下のWebページをご覧ください"
    usage_info = tk.Label(root, text=usage_text, justify="left", font=("Yu Gothic", 12))
    usage_info.pack(pady=10)

    # ホームページのリンク
    link = tk.Label(root, text="URLはこちら", fg="blue", cursor="hand2", font=("Yu Gothic", 12, "underline"))
    link.pack(pady=10)
    link.bind("<Button-1>", lambda e: webbrowser.open("https://abiding-delivery-6d9.notion.site/1264d14a044c804f9dc7e41ce20a920f"))  # ここに実際のURLを入れてください

    usage_text = "爆速議事録をご利用いただきありがとうございます！"
    usage_info = tk.Label(root, text=usage_text, justify="left", font=("Yu Gothic", 12))
    usage_info.pack(pady=10)

    # 問い合わせのリンク
    contact_link = tk.Label(root, text="お問い合わせ、バグの報告はこちらのアカウントまで", fg="blue", cursor="hand2", font=("Yu Gothic", 12, "underline"))
    contact_link.pack(pady=10)
    contact_link.bind("<Button-1>", lambda e: webbrowser.open("https://x.com/petit_hiroto"))  # ここに実際のURLを入れてください

    # 戻るボタンを右上に配置
    back_button = tk.Button(root, text="戻る", command=show_main_menu, width=5, height=1)
    back_button.place(x=800, y=20)
    back_button.lift()  # ボタンを最前面に配置

def save_prompt_to_settings(prompt_text):
    """プロンプトをsettings.jsonに保存する関数"""
    ensure_settings_exist()  # フォルダとファイルの存在を確認
    settings_path = get_settings_path()
    try:
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
        settings['transcription_prompt'] = prompt_text
        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        logging.info("プロンプトがsettings.jsonに保存されました。")
        messagebox.showinfo("保存", "プロンプトが保存されました。")
    except Exception as e:
        logging.error(f"プロンプトの保存中にエラーが発生しました: {str(e)}")
        messagebox.showerror("エラー", "プロンプトの保存中にエラーが発生しました。")

def load_output_directory():
    """settings.jsonから出力先ディレクトリを読み込む関数"""
    settings_path = get_settings_path()
    if os.path.exists(settings_path):
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
            return settings.get('output_directory', '')  # デフォルト値を空文字に変更
    return ''  # ファイルが存在しない場合も空文字を返す

def save_output_directory_to_settings(directory):
    """出力先ディレクトリをsettings.jsonに保存する関数"""
    ensure_settings_exist()  # フォルダとファイルの存在を確認
    settings_path = get_settings_path()
    try:
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
        settings['output_directory'] = directory
        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        logging.info("出力先ディレクトリがsettings.jsonに保存されました。")
    except Exception as e:
        logging.error(f"出力先ディレクトリの保存中にエラーが発生しました: {str(e)}")

def load_api_keys():
    """settings.jsonからAPIキーを読み込む関数"""
    settings_path = get_settings_path()  # 修正: get_current_dir() から get_settings_path() に変更
    if os.path.exists(settings_path):
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
            return [settings['gemini_api_keys'][f'GEMINI_API_KEY_{i}'] for i in range(1, 11)]
    logging.error("settings.jsonが見つからないか、APIキーが設定されていません。")
    return []

def get_api_keys_text():
    """APIキーをテキストボックスに表示するための文字列を生成する関数"""
    api_keys = load_api_keys()
    return "\n".join(api_keys)

def save_api_keys_to_settings(api_keys_text):
    """APIキーをsettings.jsonに保存する関数"""
    ensure_settings_exist()  # フォルダとファイルの存在を確認
    settings_path = get_settings_path()
    try:
        with open(settings_path, 'r', encoding='utf-8') as f:
            settings = json.load(f)
        api_keys = api_keys_text.strip().split('\n')
        settings['gemini_api_keys'] = {f'GEMINI_API_KEY_{i+1}': key for i, key in enumerate(api_keys)}
        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        logging.info("APIキーがsettings.jsonに保存されました。")
        messagebox.showinfo("保存", "APIキーが保存されました。")
    except Exception as e:
        logging.error(f"APIキーの保存中にエラーが発生しました: {str(e)}")
        messagebox.showerror("エラー", "APIキーの保存中にエラーが発生しました。")

def get_settings_path():
    # ユーザーディレクトリのアプリケーションデータフォルダに保存
    return Path.home() / ".my_app" / "settings.json"

def load_settings():
    settings_path = get_settings_path()
    if settings_path.exists():
        with open(settings_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_settings(settings=None):
    try:
        settings_path = get_settings_path()
        settings_path.parent.mkdir(parents=True, exist_ok=True)  # フォルダが存在しない場合は作成

        # settings.jsonが存在しない場合、デフォルトの設定を作成
        if not settings_path.exists():
            settings = {
                'transcription_prompt': '',
                'output_directory': str(Path.home() / 'Documents'),
                'gemini_api_keys': {f'GEMINI_API_KEY_{i+1}': '' for i in range(10)}
            }

        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        print("Settings saved successfully.")  # ログ出力
    except Exception as e:
        print(f"Error saving settings: {e}")  # エラーログ

def ensure_settings_exist():
    settings_path = Path.home() / ".my_app" / "settings.json"
    
    # フォルダが存在しない場合は作成
    if not settings_path.parent.exists():
        settings_path.parent.mkdir(parents=True, exist_ok=True)
        print(f"フォルダを作成しました: {settings_path.parent}")
    
    # settings.jsonが存在しない場合は作成
    if not settings_path.exists():
        settings = {
            'transcription_prompt': '',
            'output_directory': '',
            'gemini_api_keys': {f'GEMINI_API_KEY_{i+1}': '' for i in range(10)}
        }
        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
        print(f"settings.jsonを作成しました: {settings_path}")
    else:
        print(f"settings.jsonは既に存在します: {settings_path}")

# 確認と作成を実行
ensure_settings_exist()


def complete_audio_upload():
    global processing_done, start_time, selected_file_name, estimated_time_text
    if selected_file:
        processing_done = False  # この行を追加
        start_time = time.time()  # 処理開始時刻を記録
        selected_file_name = os.path.basename(selected_file)
        file_size_mb = os.path.getsize(selected_file) / (1024 * 1024)  # MBに変換

        # 想定処理時間を計算
        if file_size_mb <= 10:
            estimated_time = "0~1分"
        elif file_size_mb <= 20:
            estimated_time = "1〜2分"
        else:
            estimated_time = "1〜2分"

        estimated_time_text = estimated_time  # 時間だけを保存
        estimated_time_label.config(text=f"想定処理時間：{estimated_time}")  # 表示時にテキストを追加

        root.update_idletasks()
        processed_files = load_processed_files()
        threading.Thread(target=process_audio_file_async, args=(selected_file, processed_files, start_time)).start()
    else:
        messagebox.showwarning("警告", "ファイルが選択されていません。")

def upload_audio_file():
    global selected_file, selected_file_name, estimated_time_text
    selected_file = filedialog.askopenfilename(filetypes=[("Audio Files", "*.wav *.mp3 *.m4a")])
    if selected_file:
        selected_file_name = os.path.basename(selected_file)
        file_label.config(text=f"選択したファイル\n{selected_file_name}")
        
        # ファイルサイズを取得
        file_size_mb = os.path.getsize(selected_file) / (1024 * 1024)  # MBに変換
        
        # 想定処理時間を計算
        if file_size_mb <= 10:
            estimated_time = "0〜1分"
        elif file_size_mb <= 20:
            estimated_time = "1〜2分"
        else:
            estimated_time = "1〜2分"
        
        # 想定処理時間を表示（「想定処理時間：」を含める）
        estimated_time_text = estimated_time  # 時間だけを保存
        estimated_time_label.config(text=f"想定処理時間：{estimated_time}")  # 表示時にテキストを追加

def upload_xlsx_file():
    global selected_file
    selected_file = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    if selected_file:
        excel_file_label.config(text=f"選択したファイル\n{os.path.basename(selected_file)}")

def complete_xlsx_upload():
    if selected_file:
        root.update_idletasks()
        threading.Thread(target=process_xlsx_file_async, args=(selected_file,)).start()
    else:
        messagebox.showwarning("警告", "ファイルが選択されていません。")

def process_xlsx_file_async(xlsx_file):
    try:
        if getattr(sys, 'frozen', False):
            template_path = os.path.join(sys._MEIPASS, 'template.docx')
        else:
            template_path = os.path.join(get_current_dir(), 'template.docx')

        output_directory = load_output_directory()
        output_path = os.path.join(output_directory, f"{os.path.splitext(os.path.basename(xlsx_file))[0]}_議事録.docx")
        
        success = create_minutes(xlsx_file, template_path, output_path)
        if success:
            root.after(0, lambda: reset_file_info())
            root.after(0, lambda: (messagebox.showinfo("完了", "議事録の作成が完了しました。"), show_main_menu()))
        else:
            messagebox.showerror("エラー", "ファイルの処理中にエラーが発生しました。")
    except Exception as e:
        logging.error(f"Excel処理中にエラーが発生: {str(e)}")
        messagebox.showerror("エラー", f"処理中にエラーが発生しました: {str(e)}")

def process_audio_file_async(audio_file, processed_files, start_time):
    global processing_done, selected_file, selected_file_name, estimated_time_text

    def update_elapsed_time(start_time=start_time):  # nonlocalの代わりにパラメータとして渡す
        if processing_done:
            uploading_label.config(text="")  # 経過時間表示をクリア
            return

        if start_time:
            try:
                elapsed_time = int(time.time() - start_time)
                minutes, seconds = divmod(elapsed_time, 60)
                uploading_label.config(text=f"経過時間: {minutes}分{seconds}秒")
                if not processing_done:
                    root.after(1000, lambda: update_elapsed_time(start_time))
            except Exception as e:
                logging.error(f"経過時間の更新中にエラーが発生: {str(e)}")

    # メインスレッドで最初の更新を開始
    root.after(0, update_elapsed_time)

    # プロンプトが空でないか確認
    if not transcription_prompt:
        logging.error("プロンプトが空です。音声ファイルの処理を中止します。")
        root.after(0, lambda: messagebox.showerror("エラー", "プロンプトが空です。処理を中止します。"))
        return

    try:
        logging.info(f"{audio_file}の処理を開始します。")
        success = process_audio_file(audio_file, processed_files)
        processing_done = True  # 処理完了を示すために True を設定

        if success:
            # 処理が成功した場合、選択したファイル情報と想定処理時間をリセット
            root.after(0, lambda: reset_file_info())
            root.after(0, lambda: (messagebox.showinfo("完了", "ファイルのアップロードが完了しました。"), show_main_menu()))
        else:
            root.after(0, lambda: messagebox.showerror("エラー", "ファイルの処理中にエラーが発生しました。"))
    except Exception as e:
        logging.exception(f"音声ファイルの処理中にエラーが発生しました: {str(e)}")
        root.after(0, lambda: messagebox.showerror("エラー", "音声ファイルの処理中にエラーが発生しました。"))

def reset_file_info():
    global selected_file, selected_file_name, estimated_time_text, processing_done, start_time
    selected_file = None
    selected_file_name = ""
    estimated_time_text = ""
    start_time = None
    processing_done = True  # ここを False から True に変更
    # UIの表示をリセット
    file_label.config(text="選択したファイル: なし")
    excel_file_label.config(text="選択したファイル: なし")  # Excel用のラベルもリセット
    estimated_time_label.config(text="")
    uploading_label.config(text="")

def add_dll_directory():
    if getattr(sys, 'frozen', False):
        dll_path = os.path.join(sys._MEIPASS, 'DLLs')
        if os.path.exists(dll_path):
            os.add_dll_directory(dll_path)

# アプリケーション起動時に呼び出し
add_dll_directory()

if __name__ == "__main__":
    main()